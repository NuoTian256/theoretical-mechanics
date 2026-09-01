from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

doc = Document()

# Helper: set Chinese font
def set_font(run, name='宋体', size=10.5, bold=False, color=None):
    font = run.font
    font.name = name
    run._element.rPr.rFonts.set(qn('w:eastAsia'), name)
    font.size = Pt(size)
    font.bold = bold
    if color:
        font.color.rgb = color

def add_heading(text, level=1):
    p = doc.add_heading(level=level)
    run = p.add_run(text)
    if level == 1:
        set_font(run, '黑体', 16, bold=True, color=RGBColor(0x00, 0x00, 0x80))
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif level == 2:
        set_font(run, '黑体', 14, bold=True, color=RGBColor(0x80, 0x00, 0x00))
    else:
        set_font(run, '黑体', 12, bold=True)
    return p

def add_para(text, bold=False, indent=True, size=10.5):
    p = doc.add_paragraph()
    if indent:
        p.paragraph_format.first_line_indent = Inches(0.3)
    run = p.add_run(text)
    set_font(run, '宋体', size, bold=bold)
    return p

def add_bullet(text, level=0):
    p = doc.add_paragraph(style='List Bullet' if level == 0 else 'List Bullet 2')
    run = p.add_run(text)
    set_font(run, '宋体', 10.5)
    return p

# Title
add_heading('理论力学（分析力学）复习提纲', level=1)
add_para('根据《clas_mech 讲义》《然=一淼手写笔记》《普通班理论力学平时大作业》及回忆考题整理', size=9)
add_para('区分：【要背】= 必须准确记忆的定义、公式、结论；【要了解】= 理解推导思路、会套用解题流程', size=9)
doc.add_paragraph()

# Exam mapping
add_heading('一、回忆考题与考点对应', level=2)
add_para('先对照你回忆的题型，明确复习重心：', bold=True)
rows = [
    ('概念题', '完整/非完整约束、相空间、简正坐标、虚功原理、哈密顿原理、正则变换、拉格朗日方程、达朗贝尔原理'),
    ('选择题', '约束分类、循环坐标与广义动量守恒、哈密顿原理的变量（q, q̇, t 或 q, p, t）、完整约束判定'),
    ('大题 1', '耦合摆：写拉氏量 → 小振动近似 → 久期方程 → 本征频率 → 简正坐标'),
    ('大题 2', '母函数证明：四类母函数对应的正则变换公式，或从勒让德变换推导正则方程'),
    ('大题 3', '滑轮/阿特伍德机：分别用牛顿力学、拉格朗日力学、哈密顿力学求解同一题'),
    ('大题 4', '两自由度滑轮组：建立拉氏量（T−V），代入拉格朗日方程；也可尝试哈密顿量'),
    ('大题 5', '虚功原理课本例题：杆靠在光滑墙/棱角、双杆铰链受水平力等静力平衡'),
]
table = doc.add_table(rows=1, cols=2)
table.style = 'Light Grid Accent 1'
hdr = table.rows[0].cells
hdr[0].text = '题型'
hdr[1].text = '对应考点'
for r, c in rows:
    row = table.add_row().cells
    row[0].text = r
    row[1].text = c
for row in table.rows:
    for cell in row.cells:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                set_font(run, '宋体', 10)
doc.add_paragraph()

# Must memorize
add_heading('二、【要背】核心概念与公式', level=2)

add_heading('2.1 约束与自由度', level=3)
add_bullet('完整约束：几何约束 + 可积分的微分约束（最终能写成 f(r,t)=0）；非完整约束：不可积分的微分约束。本课程只研究完整系统。')
add_bullet('定常约束：不显含 t；非定常约束：显含 t。')
add_bullet('理想约束：约束力虚功之和为零，即 Σ R_i · δr_i = 0。')
add_bullet('自由度 s = 3N − k（k 为完整约束数），广义坐标数 = 自由度。')

add_heading('2.2 拉格朗日力学', level=3)
add_bullet('拉氏量：L = T − V（保守系统）。L 是 q, q̇, t 的函数。')
add_bullet('欧拉-拉格朗日方程：d/dt(∂L/∂q̇_α) − ∂L/∂q_α = 0，α=1,…,s。')
add_bullet('广义动量：p_α = ∂L/∂q̇_α。')
add_bullet('循环坐标（可遗坐标）：L 不显含 q_α，则 p_α = const。')
add_bullet('哈密顿原理（最小作用量原理）：δS = δ∫_{t1}^{t2} L dt = 0。')
add_bullet('拉氏量非唯一性：L′ = L + df/dt 给出相同运动方程。')

add_heading('2.3 哈密顿力学', level=3)
add_bullet('勒让德变换：H = Σ p_α q̇_α − L = H(q,p,t)。')
add_bullet('正则方程：q̇_α = ∂H/∂p_α，ṗ_α = −∂H/∂q_α。')
add_bullet('相空间：由 s 个 q_α 和 s 个 p_α 张成的 2s 维空间。')
add_bullet('若 H 不显含 t，则 H 守恒；定常约束下 H = E（总能量）。')
add_bullet('泊松括号：[φ,ψ] = Σ_α (∂φ/∂q_α · ∂ψ/∂p_α − ∂φ/∂p_α · ∂ψ/∂q_α)。')
add_bullet('运动方程的泊松括号形式：df/dt = ∂f/∂t + [f,H]；q̇_α = [q_α,H]，ṗ_α = [p_α,H]。')
add_bullet('基本泊松括号：[q_α,q_β]=0，[p_α,p_β]=0，[q_α,p_β]=δ_{αβ}。')

add_heading('2.4 虚功原理', level=3)
add_bullet('虚位移 δr_i：δt=0，被约束允许的假想无限小位移。')
add_bullet('虚功原理：理想约束系统平衡时，主动力虚功之和为零，Σ F_i · δr_i = 0。')
add_bullet('用广义坐标表示：δW = Σ Q_α δq_α = 0，且 Q_α = 0（q_α 独立）。')
add_bullet('保守系统平衡：δV = 0，即势能在平衡位置取极值。')

add_heading('2.5 小振动与简正坐标', level=3)
add_bullet('小振动：完整保守系统在稳定平衡位置附近的运动。')
add_bullet('拉氏量：L ≈ 1/2 Σ_{αβ} m_{αβ} q̇_α q̇_β − 1/2 Σ_{αβ} c_{αβ} q_α q_β。')
add_bullet('运动方程：Σ_β (m_{αβ} q̈_β + c_{αβ} q_β) = 0。')
add_bullet('久期方程：det(c − ω²m) = 0，解出 s 个本征频率 ω_i。')
add_bullet('简正坐标：使每个坐标以单一频率振动的新广义坐标；可通过把 m、c 同时对角化或观察振幅矩阵得到。')
add_bullet('耦合摆本征频率：ω₁ = √(g/l)，ω₂ = √(g/l + 2k/m)；简正坐标：ξ₁=(θ+φ)/2，ξ₂=(θ−φ)/2。')

add_heading('2.6 正则变换与母函数', level=3)
add_bullet('正则变换：保持正则方程形式不变的相空间变换。')
add_bullet('四类母函数及变换公式：')
add_bullet('U₁(q,Q,t)：p_α=∂U₁/∂q_α，P_α=−∂U₁/∂Q_α，H̃=H+∂U₁/∂t', level=1)
add_bullet('U₂(q,P,t)：p_α=∂U₂/∂q_α，Q_α=∂U₂/∂P_α，H̃=H+∂U₂/∂t', level=1)
add_bullet('U₃(p,Q,t)：q_α=−∂U₃/∂p_α，P_α=−∂U₃/∂Q_α，H̃=H+∂U₃/∂t', level=1)
add_bullet('U₄(p,P,t)：q_α=−∂U₄/∂p_α，Q_α=∂U₄/∂P_α，H̃=H+∂U₄/∂t', level=1)
add_bullet('哈密顿-雅可比方程：H(q, ∂S/∂q, t) + ∂S/∂t = 0；S 为作用量。')

add_heading('2.7 守恒与对称（诺特定理）', level=3)
add_bullet('诺特定理：每一种连续对称性对应一个守恒量。')
add_bullet('空间平移对称 → 动量守恒；空间旋转对称 → 角动量守恒；时间平移对称 → 能量守恒。')
add_bullet('若 L 不显含 t，能量函数 h = Σ p_α q̇_α − L 守恒；定常约束下 h=E。')

# Must understand
add_heading('三、【要了解】解题方法与推导思路', level=2)

add_heading('3.1 拉格朗日力学解题“三步走”', level=3)
add_bullet('1. 分析约束，确定自由度 s，选择广义坐标 q_α（越少越好、越简单越好）。')
add_bullet('2. 用广义坐标写出 T 和 V，得到 L = T − V。注意：T 要代回到广义坐标后再求导。')
add_bullet('3. 代入 d/dt(∂L/∂q̇_α) − ∂L/∂q_α = 0，整理得到运动方程。每个 q_α 对应一个方程。')

add_heading('3.2 哈密顿力学解题“三步走”', level=3)
add_bullet('1. 先写 L（同拉格朗日法）。')
add_bullet('2. 求广义动量 p_α = ∂L/∂q̇_α，反解 q̇_α = q̇_α(q,p,t)，构造 H = Σ p_α q̇_α − L。')
add_bullet('3. 代入正则方程 q̇_α=∂H/∂p_α，ṗ_α=−∂H/∂q_α。')

add_heading('3.3 虚功原理解题“三步走”', level=3)
add_bullet('1. 确定自由度，选广义坐标；分析主动力，忽略约束力（理想约束）。')
add_bullet('2. 写出虚功 δW = Σ F_i · δr_i，把所有量用广义坐标变分 δq_α 表示。')
add_bullet('3. 由于 δq_α 独立，令各系数为零，得到平衡条件。保守系统也可令 ∂V/∂q_α=0。')

add_heading('3.4 小振动解题“五步曲”', level=3)
add_bullet('1. 选广义坐标，取稳定平衡位置为 q=0，且 V(0)=0。')
add_bullet('2. 写 L = T − V，做二阶近似：T 为 q̇ 的二次型，V 为 q 的二次型。')
add_bullet('3. 读出质量矩阵 m_{αβ} 和弹性矩阵 c_{αβ}。')
add_bullet('4. 解久期方程 det(c − ω²m)=0，得到本征频率 ω_i。')
add_bullet('5. 代回线性方程组求振幅比，得到简正坐标。')

add_heading('3.5 正则变换解题“四步曲”', level=3)
add_bullet('1. 判断母函数类型（变量是 q,Q 还是 q,P 等）。')
add_bullet('2. 用对应公式写出 p、P（或 Q）与旧变量的关系。')
add_bullet('3. 把旧 H 用新变量表示得到 H̃，代入新正则方程。')
add_bullet('4. 解出新变量，再反解回旧坐标。')

add_heading('3.6 必须掌握的经典例题', level=3)
add_bullet('阿特伍德机/滑轮：分别用牛顿法、拉格朗日法、哈密顿法做一遍，比较异同。')
add_bullet('耦合摆：记熟 L、久期方程、两个频率、简正坐标的物理意义。')
add_bullet('双摆小振动：笔记与讲义都有，重点练写 T 和 V。')
add_bullet('杆靠光滑墙/棱角：虚功原理典型题，注意几何关系 y_c = l sinθ − d tanθ。')
add_bullet('最速降线：会用欧拉-拉格朗日方程写出微分方程，知道解是摆线。')

add_heading('3.7 重点推导（理解即可，会背关键步骤）', level=3)
add_bullet('从达朗贝尔原理推导拉格朗日方程：关键是用 ∂r_i/∂q_α 表示虚位移，并用到 d/dt(∂r_i/∂q_α)=∂v_i/∂q_α。')
add_bullet('勒让德变换推导正则方程：H=Σp q̇−L，两边求微分，对照系数即得。')
add_bullet('母函数给出正则变换：从 δ∫(p q̇−H)dt=0 出发，被积函数相差 dU/dt，比较全微分即得四类公式。')
add_bullet('泊松定理：若 f,g 为运动积分，则 [f,g] 也是（用雅可比恒等式）。')

# Exam tips
add_heading('四、考场策略与易错点', level=2)
add_bullet('概念题：完整约束强调“可写成 f(r,t)=0”；非完整约束强调“含速度且不可积分”。')
add_bullet('选择题：若 L 不含某 q_α，立刻想到 p_α 守恒；哈密顿原理的独立变分变量是 q_α（不是 p）。')
add_bullet('大题：先写“自由度 s=…，广义坐标选…”拿步骤分。')
add_bullet('滑轮题：注意绳长约束关系，选择位移/角度时要保证独立。')
add_bullet('小振动：一定要把 cosθ 近似到 1−θ²/2，弹簧伸长用水平位移近似 l(θ−φ)。')
add_bullet('虚功原理：δy 要会正确变分，常出现 cosθ、tanθ、sec²θ 的系数。')
add_bullet('母函数题：先看函数变量，对应四类公式，不要混淆正负号。')

# Footer
add_para('祝你复习顺利，考试高分！', bold=True, indent=False, size=11)

# Save
out_path = 'D:/desktop/learn/物理/理论力学/理论力学复习提纲.docx'
doc.save(out_path)
print(f'Saved: {out_path}')
